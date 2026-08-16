"""
Extract text from Word documents (.docx).

See extract_pdf_text.py for the PDF equivalent and extract_image_text.py
for standalone images — extract_documents.py is the orchestrator that
dispatches to all three based on detect_format.py's classification.
Previously DOCX files were classified but never actually extracted
(audit finding #5).
"""
from pathlib import Path
from datetime import datetime, timezone

import docx

from config import logger


def extract_docx_text(docx_path):
    """
    Extract text from a Word document.
    Returns: dict with text content and metadata, same shape as
    extract_pdf_text.extract_pdf_text() so both flow through the same
    indexing path unmodified.
    """
    docx_path = Path(docx_path)

    try:
        document = docx.Document(docx_path)

        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

        # Technical manuals put real content in tables (parts lists, spec
        # sheets, torque values) — paragraphs alone would silently drop
        # all of it, undermining the whole point of indexing these files.
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))

        full_text = "\n".join(paragraphs)
        core_props = document.core_properties

        return {
            'status': 'success',
            'filename': docx_path.name,
            'filepath': str(docx_path),
            'file_size': docx_path.stat().st_size,
            'num_pages': None,  # DOCX has no fixed page count outside a layout renderer
            'metadata': {
                'title': core_props.title or 'N/A',
                'author': core_props.author or 'N/A',
                'subject': core_props.subject or 'N/A',
            },
            'text': full_text,
            'text_length': len(full_text),
            'ocr_pages_used': [],
            'extracted_at': datetime.now(timezone.utc).isoformat(),
        }
    except Exception as e:
        logger.exception("Failed to extract %s", docx_path)
        return {
            'status': 'error',
            'filename': docx_path.name,
            'filepath': str(docx_path),
            'error': str(e),
        }
