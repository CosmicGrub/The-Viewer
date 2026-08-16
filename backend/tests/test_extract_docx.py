"""
Tests for extract_docx_text.py (finding #5 — DOCX files were classified by
detect_format.py but never actually extracted before this).
"""
import docx

from extract_docx_text import extract_docx_text


def _make_docx(tmp_path, paragraphs=(), table_rows=None, title=None, author=None):
    document = docx.Document()
    if title:
        document.core_properties.title = title
    if author:
        document.core_properties.author = author
    for text in paragraphs:
        document.add_paragraph(text)
    if table_rows:
        table = document.add_table(rows=0, cols=len(table_rows[0]))
        for row_values in table_rows:
            row = table.add_row()
            for cell, value in zip(row.cells, row_values):
                cell.text = value
    path = tmp_path / "sample.docx"
    document.save(path)
    return path


def test_extracts_paragraph_text(tmp_path):
    path = _make_docx(tmp_path, paragraphs=["Torque the fitting to 45 ft-lb.", "Replace the gasket."])

    result = extract_docx_text(path)

    assert result["status"] == "success"
    assert "Torque the fitting to 45 ft-lb." in result["text"]
    assert "Replace the gasket." in result["text"]
    assert result["text_length"] == len(result["text"])


def test_extracts_table_cell_text(tmp_path):
    # Technical manuals put real content (parts lists, spec values) in
    # tables — paragraph-only extraction would silently drop it.
    path = _make_docx(tmp_path, table_rows=[["Part No.", "5310-01-234-5678"], ["Qty", "4"]])

    result = extract_docx_text(path)

    assert result["status"] == "success"
    assert "5310-01-234-5678" in result["text"]


def test_extracts_core_metadata(tmp_path):
    path = _make_docx(tmp_path, paragraphs=["body"], title="Hydraulic Pump Manual", author="Dept. of the Army")

    result = extract_docx_text(path)

    assert result["metadata"]["title"] == "Hydraulic Pump Manual"
    assert result["metadata"]["author"] == "Dept. of the Army"


def test_missing_file_returns_error_status(tmp_path):
    result = extract_docx_text(tmp_path / "does_not_exist.docx")

    assert result["status"] == "error"
    assert "error" in result
